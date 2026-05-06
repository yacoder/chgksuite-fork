#!/usr/bin/env python
#! -*- coding: utf-8 -*-
import contextlib
import inspect
import json
import os
import random
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from io import BytesIO

import pytest
import chgksuite.parser as parser_module
from chgksuite.common import (
    DefaultArgs,
    image_data_to_jpeg_bytes,
    optimize_raster_image_data,
    read_text_file,
)
from chgksuite.composer.chgksuite_parser import parse_4s, replace_counters
from chgksuite.composer.composer_common import (
    _parse_4s_elem,
    game_to_ext,
    parseimg,
    remove_accents_standalone,
)
from chgksuite.composer.docx import (
    DocxExporter,
    add_hyperlink_to_docx,
    add_text_run_to_docx,
    optimize_docx_images,
    remove_square_brackets_standalone,
)
from chgksuite.composer.telegram import TelegramExporter
from chgksuite.parser import (
    chgk_parse_docx,
    chgk_parse_txt,
    compose_4s,
    si_parse_docx,
    si_parse_text,
    troika_parse_docx,
    troika_parse_text,
)
from chgksuite.typotools import cyr_lat_check_word, get_quotes_right, replace_no_break
from PIL import Image

currentdir = os.path.dirname(os.path.abspath(inspect.getfile(inspect.currentframe())))
parentdir = os.path.dirname(currentdir)

# Encrypted test file support
PASSWORD_FILE = os.path.join(currentdir, "tests_password.txt")


def get_test_password():
    """Read password from file, return None if not found."""
    if os.path.exists(PASSWORD_FILE):
        with open(PASSWORD_FILE, "r") as f:
            return f.read().strip()
    return None


def decrypt_test_file(filepath: str, password: str) -> bytes:
    """Decrypt a test file using XOR."""
    import hashlib

    key = hashlib.sha256(password.encode()).digest()
    with open(filepath, "rb") as f:
        data = f.read()
    key_len = len(key)
    return bytes(b ^ key[i % key_len] for i, b in enumerate(data))


with open(os.path.join(currentdir, "settings.json")) as f:
    settings = json.loads(f.read())


ljlogin, ljpassword = open(os.path.join(currentdir, "ljcredentials")).read().split("\t")


def workaround_chgk_parse(filename, game=None, **kwargs):
    if game in ("si", "troika"):
        args = DefaultArgs(**kwargs)
        if (
            not getattr(args, "numbers_handling", None)
            or args.numbers_handling == "default"
        ):
            args.numbers_handling = "all"
        if filename.endswith(".txt"):
            text = read_text_file(filename)
            if game == "si":
                return si_parse_text(text, args=args)
            return troika_parse_text(text, args=args)
        elif filename.endswith(".docx"):
            if game == "si":
                return si_parse_docx(filename, args=args)
            return troika_parse_docx(filename, args=args)
    if filename.endswith(".txt"):
        return chgk_parse_txt(filename, args=DefaultArgs(**kwargs))
    elif filename.endswith(".docx"):
        return chgk_parse_docx(filename, args=DefaultArgs(**kwargs))
    return


QUOTE_TEST_CASES = [
    ('«"Альфа" Бета»', "«„Альфа“ Бета»"),
    ("«“Альфа” Бета»", "«„Альфа“ Бета»"),
    ("«„Альфа“ Бета»", "«„Альфа“ Бета»"),
    ("«Альфа», “Бета”", "«Альфа», «Бета»"),
    (
        '"Он сказал: "Привет!". А потом заплакал"',
        "«Он сказал: „Привет!“. А потом заплакал»",
    ),
    (
        "“Он сказал: “Привет!”. А потом заплакал”",
        "«Он сказал: „Привет!“. А потом заплакал»",
    ),
    (
        "Все вопросы тура написаны по одному источнику — книге Натальи Эдуардовны Манусаджян «Применение соматопсихотерапии во время тренировок по „Что? Где? Когда?“ как метода развития креативности мышления».",
        "Все вопросы тура написаны по одному источнику — книге Натальи Эдуардовны Манусаджян «Применение соматопсихотерапии во время тренировок по „Что? Где? Когда?“ как метода развития креативности мышления».",
    ),
]


@pytest.mark.parametrize("a,expected", QUOTE_TEST_CASES)
def test_quotes(a, expected):
    assert get_quotes_right(a) == expected


def test_replace_no_break_preserves_url_hyphens():
    url = "https://example.com/history/160517-nilometer-discovered/ик-с"
    result = replace_no_break(f"Ссылка на {url}; код И-К-Б-С.")

    assert url in result
    assert f"на\u00a0{url}" in result
    assert "И\u2011К\u2011Б\u2011С" in result


# Test cases for Latin accented character conversion to Cyrillic
# Format: (input, expected_output)
# The fix ensures uppercase Cyrillic neighbors are recognized correctly
CYR_LAT_ACCENT_TEST_CASES = [
    # Bug case: Latin á (U+00E1) after uppercase Cyrillic should convert
    ("Хáральд", "Ха́ральд"),  # Х is uppercase Cyrillic
    # Latin à (U+00E0) after lowercase Cyrillic - already worked
    ("Ивàново", "Ива́ново"),
    # Latin é (U+00E9) in middle of word - already worked
    ("крылéц", "крыле́ц"),
    # Latin ó (U+00F3) after uppercase Cyrillic
    ("Óльга", "О́льга"),  # О is uppercase Cyrillic
    # Latin ú (U+00FA) mapped to Cyrillic и́
    ("Иúсус", "Ии́сус"),
    # Multiple accented chars in one word
    ("Москвá", "Москва́"),
    # Pure Latin word - should NOT convert (no Cyrillic neighbors)
    ("café", None),  # None means no change
    # Mixed but Latin char surrounded by Latin - should NOT convert
    ("Caféшоп", None),  # é surrounded by Latin 'f' and Cyrillic 'ш', but 'f' blocks it
    # Single char word - should not convert (length check)
    ("á", None),
    # Uppercase Latin accent after uppercase Cyrillic
    ("ХÁРАЛЬД", "ХА́РАЛЬД"),
]


@pytest.mark.parametrize("input_word,expected", CYR_LAT_ACCENT_TEST_CASES)
def test_cyr_lat_accent_conversion(input_word, expected):
    result = cyr_lat_check_word(input_word)
    if expected is None:
        assert result is None, f"Expected no change for '{input_word}', got '{result}'"
    else:
        assert result == expected, (
            f"Expected '{expected}' for '{input_word}', got '{result}'"
        )


with open(os.path.join(parentdir, "chgksuite", "resources", "regexes_ru.json")) as f:
    TEST_REGEXES = json.load(f)


SQUARE_BRACKET_TEST_CASES = [
    ("black [блэк]", "black"),
    ("black [блэк] смотрит [looks]", "black смотрит"),
    (
        "text with [Раздаточный материал: handout] here",
        "text with [Раздаточный материал: handout] here",
    ),  # handout preserved
    ("text \\[escaped\\]", "text [escaped]"),  # escaped brackets restored
    ("simple text", "simple text"),  # no brackets
]


@pytest.mark.parametrize("input_text,expected", SQUARE_BRACKET_TEST_CASES)
def test_remove_square_brackets(input_text, expected):
    assert remove_square_brackets_standalone(input_text, TEST_REGEXES) == expected


ACCENT_TEST_CASES = [
    ("при́вет", "привет"),  # \u0301 accent removed
    ("мо́ре си́нее", "море синее"),  # multiple accents
    (
        "[Раздаточный материал: при́вет]",
        "[Раздаточный материал: при́вет]",
    ),  # accent in handout preserved
    ("simple text", "simple text"),  # no accents
]


@pytest.mark.parametrize("input_text,expected", ACCENT_TEST_CASES)
def test_remove_accents(input_text, expected):
    assert remove_accents_standalone(input_text, TEST_REGEXES) == expected


def test_troika_colon_theme_after_source():
    parsed = troika_parse_text(
        """ТРОЙКА

Тема 1. ТЕМА: ВРЕМЯ

Автор: Автор

1. Первый вопрос.

Ответ: Первый ответ.

Источник: https://example.com/one

ТЕМА: ОСКАРЫ

1. Второй вопрос.

Ответ: Второй ответ.

Источник: https://example.com/two""",
        args=DefaultArgs(game="troika"),
    )

    themes = [element[1] for element in parsed if element[0] == "theme"]
    assert themes == ["ВРЕМЯ", "ОСКАРЫ"]

    questions = [element[1] for element in parsed if element[0] == "Question"]
    assert len(questions) == 2
    assert questions[0]["source"] == "https://example.com/one"
    assert questions[1]["question"] == "Второй вопрос."


def test_troika_source_list_strips_parenthesized_numbers():
    parsed = troika_parse_text(
        """ТРОЙКА

ТЕМА: ВРЕМЯ

Автор: Автор

1. Вопрос.

Ответ: Ответ.

Источник:
1) https://example.com/one
2) https://example.com/two""",
        args=DefaultArgs(game="troika"),
    )

    questions = [element[1] for element in parsed if element[0] == "Question"]
    assert questions[0]["source"] == [
        "https://example.com/one",
        "https://example.com/two",
    ]

    rendered = compose_4s(
        parsed, args=DefaultArgs(game="troika", numbers_handling="all")
    )
    assert "- 1) https://example.com/one" not in rendered
    assert "- https://example.com/one" in rendered


def test_troika_author_gratitude_after_author_is_meta():
    parsed = troika_parse_text(
        """ТРОЙКА

Автор: Артём Горячев
Автор благодарит за тестирование хороших людей.

ТЕМА: ВРЕМЯ

Автор: Артём Горячев

1. Вопрос.

Ответ: Ответ.

Источник: https://example.com/one""",
        args=DefaultArgs(game="troika"),
    )

    assert parsed[1] == ["author", "Артём Горячев"]
    assert parsed[2] == ["meta", "Автор благодарит за тестирование хороших людей."]

    rendered = compose_4s(
        parsed, args=DefaultArgs(game="troika", numbers_handling="all")
    )
    assert "@ Артём Горячев\nАвтор благодарит" not in rendered
    assert "# Автор благодарит за тестирование хороших людей." in rendered


def test_troika_number_only_question_keeps_leading_host_note():
    parsed = troika_parse_text(
        """ТРОЙКА

ТЕМА: МОРЕ

1.
[Комментарий ведущему: не объявлять, что в вопросе есть кавычки]

В 2007 году «James E. Williams» [джеймс и уильямс] спас «Тэ Хон Дан».

Ответ: Сомали́.""",
        args=DefaultArgs(game="troika"),
    )

    questions = [element[1] for element in parsed if element[0] == "Question"]
    assert len(questions) == 1
    assert questions[0]["question"] == (
        "[Комментарий ведущему: не объявлять, что в вопросе есть кавычки]\n"
        "В 2007 году «James E. Williams» [джеймс и уильямс] спас «Тэ Хон Дан»."
    )
    assert questions[0]["answer"] == "Сомали́."

    rendered = compose_4s(
        parsed, args=DefaultArgs(game="troika", numbers_handling="all")
    )
    assert "? [Комментарий ведущему" in rendered
    assert "В 2007 году «James E. Williams»" in rendered


def test_troika_pypandoc_html_preserves_ordered_list_start_numbers(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("ТРОЙКА")
    doc.add_paragraph("ТЕМА: ВРЕМЯ")
    doc.add_paragraph("Автор: Автор")
    for num in range(1, 4):
        doc.add_paragraph(f"Вопрос {num}.", style="List Number")
        doc.add_paragraph(f"Ответ: Ответ {num}.")
        doc.add_paragraph(f"Источник: https://example.com/{num}")

    filename = tmp_path / "troika_numbering.docx"
    doc.save(filename)

    parsed = troika_parse_docx(
        str(filename),
        args=DefaultArgs(game="troika", parsing_engine="pypandoc_html"),
    )

    themes = [element[1] for element in parsed if element[0] == "theme"]
    questions = [element[1] for element in parsed if element[0] == "Question"]
    assert themes == ["ВРЕМЯ"]
    assert [question["number"] for question in questions] == ["1", "2", "3"]


def test_troika_python_docx_preserves_ordered_list_start_numbers(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("ТРОЙКА")
    doc.add_paragraph("ТЕМА: ВРЕМЯ")
    doc.add_paragraph("Автор: Автор")
    for num in range(1, 4):
        doc.add_paragraph(f"Вопрос {num}.", style="List Number")
        doc.add_paragraph(f"Ответ: Ответ {num}.")
        doc.add_paragraph(f"Источник: https://example.com/{num}")

    filename = tmp_path / "troika_numbering.docx"
    doc.save(filename)

    parsed = troika_parse_docx(
        str(filename),
        args=DefaultArgs(game="troika", parsing_engine="python_docx"),
    )

    themes = [element[1] for element in parsed if element[0] == "theme"]
    questions = [element[1] for element in parsed if element[0] == "Question"]
    assert themes == ["ВРЕМЯ"]
    assert [question["number"] for question in questions] == ["1", "2", "3"]


def test_docx_to_text_python_docx_preserves_needed_docx_attributes(
    monkeypatch, tmp_path
):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    def fail_convert_file(*args, **kwargs):
        raise AssertionError("python_docx engine should not call pypandoc")

    monkeypatch.setattr(parser_module.pypandoc, "convert_file", fail_convert_file)

    image_path = tmp_path / "pixel.png"
    Image.new("RGB", (1, 1), (255, 0, 0)).save(image_path)

    doc = Document()
    doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    doc.add_heading("Раунд", level=1)
    formatted = doc.add_paragraph()
    formatted.add_run("До ")
    formatted.add_run("жирный").bold = True
    formatted.add_run(" и ")
    formatted.add_run("курсив").italic = True
    formatted.add_run(" и ")
    formatted.add_run("подчеркнутый").underline = True

    linked = doc.add_paragraph("Ссылка: ")
    add_hyperlink_to_docx(
        doc, linked, "пример", "https://example.com/path_with_under"
    )
    linked_url = doc.add_paragraph("Прямая ссылка: ")
    add_hyperlink_to_docx(
        doc,
        linked_url,
        "https://example.com/path_with_under",
        "https://example.com/path_with_under",
    )

    for num in range(1, 3):
        doc.add_paragraph(f"Вопрос {num}.", style="List Number")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"

    image_paragraph = doc.add_paragraph("Картинка ")
    image_paragraph.add_run().add_picture(str(image_path))

    filename = tmp_path / "sample.docx"
    doc.save(filename)

    text = parser_module.docx_to_text(
        str(filename),
        args=DefaultArgs(parsing_engine="python_docx", preserve_formatting=True),
        inject_heading_markers=True,
    )

    assert "$$H1$$ Раунд" in text
    assert "До __жирный__ и _курсив_ и ___подчеркнутый___" in text
    assert "Ссылка: пример (https://example.com/path_with_under)" in text
    assert "Прямая ссылка: https://example.com/path_with_under" in text
    assert r"path\_with\_under" not in text
    assert "1. Вопрос 1." in text
    assert "2. Вопрос 2." in text
    assert "| A | B |" in text
    assert "| 1 | 2 |" in text
    assert "(img sample_001.png)" in text
    assert (tmp_path / "sample_001.png").exists()


def test_chgk_parse_txt_keeps_url_underscores_unescaped(tmp_path):
    filename = tmp_path / "sample.txt"
    filename.write_text(
        "1. Вопрос\nОтвет: Ответ\n"
        "Источник: file_name https://example.com/path_with_under\n",
        encoding="utf-8",
    )

    parsed = chgk_parse_txt(str(filename), encoding="utf-8", args=DefaultArgs())
    question = [element[1] for element in parsed if element[0] == "Question"][0]

    assert question["source"] == r"file\_name https://example.com/path_with_under"


def test_docx_to_text_pypandoc_keeps_url_underscores_unescaped(
    monkeypatch, tmp_path
):
    def fake_convert_file(*args, **kwargs):
        return "Text snake_case https://example.com/path_with_under"

    monkeypatch.setattr(parser_module.pypandoc, "convert_file", fake_convert_file)

    text = parser_module.docx_to_text(
        str(tmp_path / "test.docx"),
        args=DefaultArgs(parsing_engine="pypandoc"),
    )

    assert text == r"Text snake\_case https://example.com/path_with_under"


def test_docx_to_text_html_keeps_href_url_underscores_unescaped(
    monkeypatch, tmp_path
):
    def fake_convert_file(*args, **kwargs):
        return (
            '<p>Text snake_case <a href="https://example.com/path_with_under">'
            "label</a></p>"
        )

    monkeypatch.setattr(parser_module.pypandoc, "convert_file", fake_convert_file)

    text = parser_module.docx_to_text(
        str(tmp_path / "test.docx"),
        args=DefaultArgs(parsing_engine="pypandoc_html"),
    )

    assert r"Text snake\_case label (https://example.com/path_with_under)" in text


@pytest.mark.parametrize(
    ("parsing_engine", "pandoc_format", "converted"),
    [
        ("pypandoc", "plain", "Первый вопрос.\n\nОтвет: Первый ответ."),
        (
            "pypandoc_html",
            "html",
            "<p>Первый вопрос.</p><p>Ответ: Первый ответ.</p>",
        ),
    ],
)
def test_docx_to_text_installs_pandoc_when_missing(
    monkeypatch, capsys, tmp_path, parsing_engine, pandoc_format, converted
):
    attempts = []
    installed = []

    def fake_convert_file(source_file, to, **kwargs):
        attempts.append((source_file, to, kwargs))
        if len(attempts) == 1:
            raise OSError("No pandoc was found: either install pandoc")
        return converted

    def fake_install_pandoc():
        installed.append(True)

    monkeypatch.setattr(parser_module.pypandoc, "convert_file", fake_convert_file)
    monkeypatch.setattr(
        parser_module.pypandoc, "install_pandoc", fake_install_pandoc, raising=False
    )

    text = parser_module.docx_to_text(
        str(tmp_path / "test.docx"),
        args=DefaultArgs(parsing_engine=parsing_engine),
    )

    assert "Первый вопрос." in text
    assert "Ответ: Первый ответ." in text
    assert installed == [True]
    assert len(attempts) == 2
    assert attempts[0][1] == pandoc_format
    assert capsys.readouterr().out == "pandoc not found, installing...\n"


def test_docx_to_text_does_not_install_pandoc_for_other_pypandoc_errors(
    monkeypatch, capsys, tmp_path
):
    installed = []

    def fake_convert_file(*args, **kwargs):
        raise OSError("permission denied")

    def fake_install_pandoc():
        installed.append(True)

    monkeypatch.setattr(parser_module.pypandoc, "convert_file", fake_convert_file)
    monkeypatch.setattr(
        parser_module.pypandoc, "install_pandoc", fake_install_pandoc, raising=False
    )

    with pytest.raises(OSError, match="permission denied"):
        parser_module.docx_to_text(
            str(tmp_path / "test.docx"),
            args=DefaultArgs(parsing_engine="pypandoc"),
        )

    assert installed == []
    assert capsys.readouterr().out == ""


@contextlib.contextmanager
def make_temp_directory(**kwargs):
    temp_dir = tempfile.mkdtemp(**kwargs)
    yield temp_dir
    shutil.rmtree(os.path.abspath(temp_dir))


def normalize(string):
    return string.replace("\r\n", "\n")


def source_filename_from_canon(filename):
    if filename.endswith(".encrypted.canon"):
        return filename[: -len(".encrypted.canon")]
    return filename[: -len(".canon")]


def read_canon_text(filename):
    if filename.endswith(".encrypted.canon"):
        password = get_test_password()
        if password is None:
            pytest.skip("No password file found for encrypted test")
        return decrypt_test_file(os.path.join(currentdir, filename), password).decode(
            "utf-8"
        )
    with open(os.path.join(currentdir, filename), "r", encoding="utf-8") as f:
        return f.read()


def canon_compose_args(game):
    args = DefaultArgs(game=game)
    if game in ("si", "troika"):
        args.numbers_handling = "all"
    return args


# Regular canon files (always run)
CANON_FILENAMES = [
    fn
    for fn in os.listdir(currentdir)
    if fn.endswith(".canon") and not fn.endswith(".encrypted.canon")
]

# Add encrypted canon files only if password exists
if os.path.exists(PASSWORD_FILE):
    CANON_FILENAMES.extend(
        [fn for fn in os.listdir(currentdir) if fn.endswith(".encrypted.canon")]
    )


@pytest.mark.parametrize("filename", CANON_FILENAMES)
def test_canonical_equality(parsing_engine, filename):
    # Handle encrypted files
    is_encrypted = filename.endswith(".encrypted.canon")
    if is_encrypted:
        password = get_test_password()
        if password is None:
            pytest.skip("No password file found for encrypted test")

    with make_temp_directory(dir=".") as temp_dir:
        if is_encrypted:
            # filename = "file.docx.encrypted.canon" (16 chars for ".encrypted.canon")
            # Decrypt .encrypted.canon -> .canon in temp dir
            canon_content = decrypt_test_file(
                os.path.join(currentdir, filename), password
            )
            decrypted_canon = filename[:-16] + ".canon"  # "file.docx.canon"
            with open(os.path.join(temp_dir, decrypted_canon), "wb") as f:
                f.write(canon_content)

            # Decrypt source file (.docx.encrypted)
            source_encrypted = filename[:-6]  # remove ".canon" -> "file.docx.encrypted"
            source_decrypted = filename[
                :-16
            ]  # remove ".encrypted.canon" -> "file.docx"
            source_content = decrypt_test_file(
                os.path.join(currentdir, source_encrypted), password
            )
            with open(os.path.join(temp_dir, source_decrypted), "wb") as f:
                f.write(source_content)

            to_parse_fn = source_decrypted
            canon_fn = decrypted_canon
        else:
            # Original logic for non-encrypted files
            to_parse_fn = filename[:-6]
            canon_fn = filename
            shutil.copy(os.path.join(currentdir, filename), temp_dir)
            shutil.copy(os.path.join(currentdir, to_parse_fn), temp_dir)

        print("Testing {}...".format(to_parse_fn))
        bn, _ = os.path.splitext(to_parse_fn)
        file_settings = settings.get(to_parse_fn, {})
        game = file_settings.get("game")
        call_args = [
            sys.executable,
            "-m",
            "chgksuite",
            "parse",
            "--parsing_engine",
            parsing_engine,
        ]
        if game:
            call_args.extend(["--game", game])
        call_args.append(os.path.join(temp_dir, to_parse_fn))
        if file_settings.get("cmdline_args"):
            call_args.extend(file_settings["cmdline_args"])
        subprocess.call(call_args, timeout=5, cwd=parentdir)
        out_ext = game_to_ext(game)
        with open(
            os.path.join(temp_dir, bn + "." + out_ext), "r", encoding="utf-8"
        ) as f:
            parsed = f.read()
        with open(os.path.join(temp_dir, canon_fn), "r", encoding="utf-8") as f:
            canonical = f.read()
        assert normalize(canonical) == normalize(parsed)


@pytest.mark.parametrize("filename", CANON_FILENAMES)
def test_canon_parse_compose_parse_idempotence(filename):
    source_filename = source_filename_from_canon(filename)
    game = settings.get(source_filename, {}).get("game")
    canonical = read_canon_text(filename)

    canon_structure = parse_4s(canonical, game=game)
    composed = compose_4s(canon_structure, args=canon_compose_args(game))
    reparsed_structure = parse_4s(composed, game=game)

    assert reparsed_structure == canon_structure


TO_DOCX_FILENAMES = [
    fn for fn in os.listdir(currentdir) if fn.endswith((".docx", ".txt"))
]
TO_DOCX_FILENAMES.remove("balt09-1.txt")  # TODO: rm this line once dns is fixed


@pytest.mark.parametrize("filename", TO_DOCX_FILENAMES)
def test_docx_composition(filename):
    print("Testing {}...".format(filename))
    with make_temp_directory(dir=".") as temp_dir:
        shutil.copy(os.path.join(currentdir, filename), temp_dir)
        temp_dir_filename = os.path.join(temp_dir, filename)
        game = settings.get(filename, {}).get("game")
        parsed = workaround_chgk_parse(temp_dir_filename, game=game)
        file4s = os.path.splitext(filename)[0] + "." + game_to_ext(game)
        composed_abspath = os.path.join(temp_dir, file4s)
        print(composed_abspath)
        with open(composed_abspath, "w", encoding="utf-8") as f:
            f.write(compose_4s(parsed, args=DefaultArgs(game=game)))
        call_args = [
            sys.executable,
            "-m",
            "chgksuite",
            "compose",
            "docx",
            composed_abspath,
        ]
        code = subprocess.call(call_args, timeout=5, cwd=parentdir)
        assert 0 == code


@pytest.mark.tex
def test_tex_composition():
    for filename in os.listdir(currentdir):
        if (
            filename.endswith((".docx", ".txt"))
            and filename == "Kubok_knyagini_Olgi-2015.docx"
        ):
            print("Testing {}...".format(filename))
            with make_temp_directory(dir=".") as temp_dir:
                shutil.copy(os.path.join(currentdir, filename), temp_dir)
                temp_dir_filename = os.path.join(temp_dir, filename)
                parsed = workaround_chgk_parse(temp_dir_filename)
                file4s = os.path.splitext(filename)[0] + ".4s"
                composed_abspath = os.path.join(temp_dir, file4s)
                print(composed_abspath)
                with open(composed_abspath, "w", encoding="utf-8") as f:
                    f.write(compose_4s(parsed, args=DefaultArgs()))
                code = subprocess.call(
                    [
                        sys.executable,
                        "-m",
                        "chgksuite",
                        "compose",
                        "tex",
                        composed_abspath,
                    ],
                    cwd=parentdir,
                )
                assert 0 == code


TEST_INLINE_IMAGE = """\
? какой-то Тест вопроса с (img inline test.jpg) инлайн картинкой.
! какой-то ответ
/ какой-то комментарий
^ какой-то источник
@ какой-то автор"""


def test_inline_image():
    structure = parse_4s(TEST_INLINE_IMAGE)
    question = structure[0][1]["question"]
    question_parsed = _parse_4s_elem(question)
    img = [x for x in question_parsed if x[0] == "img"]
    assert len(img) == 1
    with make_temp_directory(dir=".") as temp_dir:
        shutil.copy(os.path.join(currentdir, "test.jpg"), temp_dir)
        img_parsed = parseimg(img[0][1], tmp_dir=temp_dir)
    assert img_parsed["inline"]
    assert os.path.basename(img_parsed["imgfile"]) == "test.jpg"
    assert compose_4s(structure, DefaultArgs()).strip() == TEST_INLINE_IMAGE.strip()


def test_parse_4s_elem_does_not_parse_url_underscores_as_italic():
    url = "https://ru.wikipedia.org/wiki/Пугачёв,_Емельян_Иванович"
    parsed = _parse_4s_elem(
        f"before https://example.com/path_with_under {url} after _italic_"
    )

    assert ["hyperlink", "https://example.com/path_with_under"] in parsed
    assert ["hyperlink", url] in parsed
    assert [run for run in parsed if run[0] == "italic"] == [["italic", "italic"]]


def test_docx_hyperlink_targets_percent_encode_non_ascii_url(tmp_path):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    url = "https://ru.wikipedia.org/wiki/Абвгде_ёжзийкл_мнопрстуф_«Хцчшщы_—_эюяабв»_гдежзий_клмнопрст_уфхцчшщ"

    doc = Document()
    doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    paragraph = doc.add_paragraph()
    add_hyperlink_to_docx(doc, paragraph, url, url)
    filename = tmp_path / "link.docx"
    doc.save(filename)

    with zipfile.ZipFile(filename) as docx_file:
        rels = docx_file.read("word/_rels/document.xml.rels").decode("utf-8")
        document = docx_file.read("word/document.xml").decode("utf-8")

    target = re.search(r'Type="[^"]+/hyperlink" Target="([^"]+)"', rels).group(1)
    assert target.startswith("https://ru.wikipedia.org/wiki/%D0%90")
    assert "%C2%AB%D0%A5%D1%86%D1%87%D1%88%D1%89%D1%8B" in target
    assert "%E2%80%94" in target
    assert not re.search(r"[А-Яа-яЁё«»—]", target)
    assert url in document


def test_docx_non_breaking_hyphen_uses_word_joiners(tmp_path):
    from docx import Document

    doc = Document()
    paragraph = doc.add_paragraph()
    add_text_run_to_docx(paragraph, "В 50\u2011е годы")
    filename = tmp_path / "nbh.docx"
    doc.save(filename)

    with zipfile.ZipFile(filename) as docx_file:
        document = docx_file.read("word/document.xml").decode("utf-8")

    assert "<w:noBreakHyphen/>" not in document
    assert "\u2011" not in document
    assert "В 50\u2060-\u2060е годы" in document


def test_optimize_docx_images_recompresses_png_as_jpeg(tmp_path):
    from docx import Document

    image_path = tmp_path / "source.png"
    rng = random.Random(0)
    image = Image.new("RGB", (180, 180))
    image.putdata(
        [
            (rng.randrange(256), rng.randrange(256), rng.randrange(256))
            for _ in range(180 * 180)
        ]
    )
    image.save(image_path)
    docx_path = tmp_path / "image.docx"
    doc = Document()
    doc.add_picture(str(image_path))
    doc.save(docx_path)

    renamed_parts = optimize_docx_images(docx_path, quality=80)

    with zipfile.ZipFile(docx_path) as docx_file:
        names = docx_file.namelist()
        content_types = docx_file.read("[Content_Types].xml").decode("utf-8")
        rels = docx_file.read("word/_rels/document.xml.rels").decode("utf-8")
        image_names = [name for name in names if name.startswith("word/media/")]
        image_data = docx_file.read(image_names[0])

    assert renamed_parts == {"word/media/image1.png": "word/media/image1.jpg"}
    assert image_names == ["word/media/image1.jpg"]
    assert image_data.startswith(b"\xff\xd8")
    assert 'Extension="jpg" ContentType="image/jpeg"' in content_types
    assert 'Target="media/image1.jpg"' in rels


def test_optimize_docx_images_preserves_transparent_png(tmp_path):
    from docx import Document

    image_path = tmp_path / "transparent.png"
    Image.new("RGBA", (120, 120), (255, 0, 0, 128)).save(
        image_path, format="PNG", compress_level=0
    )
    docx_path = tmp_path / "transparent.docx"
    doc = Document()
    doc.add_picture(str(image_path))
    doc.save(docx_path)

    optimized_parts = optimize_docx_images(docx_path, quality=80)

    with zipfile.ZipFile(docx_path) as docx_file:
        names = docx_file.namelist()
        rels = docx_file.read("word/_rels/document.xml.rels").decode("utf-8")
        image_names = [name for name in names if name.startswith("word/media/")]
        image_data = docx_file.read(image_names[0])

    assert optimized_parts == {"word/media/image1.png": "word/media/image1.png"}
    assert image_names == ["word/media/image1.png"]
    assert not image_data.startswith(b"\xff\xd8")
    assert 'Target="media/image1.png"' in rels
    with Image.open(BytesIO(image_data)) as image:
        assert image.convert("RGBA").getchannel("A").getextrema()[0] < 255


def test_image_data_to_jpeg_bytes_flattens_alpha_to_white():
    source = BytesIO()
    Image.new("RGBA", (1, 1), (255, 0, 0, 128)).save(source, format="PNG")

    jpeg_data = image_data_to_jpeg_bytes(source.getvalue(), quality=90)

    assert jpeg_data is not None
    assert jpeg_data.startswith(b"\xff\xd8")
    with Image.open(BytesIO(jpeg_data)) as image:
        assert image.mode == "RGB"


def test_optimize_raster_image_data_preserves_transparent_png():
    source = BytesIO()
    image = Image.new("RGBA", (64, 64), (255, 0, 0, 128))
    image.save(source, format="PNG", compress_level=0)

    optimized = optimize_raster_image_data(
        source.getvalue(), original_extension="png", quality=80
    )

    assert optimized is not None
    extension, content_type, optimized_data = optimized
    assert extension == "png"
    assert content_type == "image/png"
    with Image.open(BytesIO(optimized_data)) as optimized_image:
        assert optimized_image.convert("RGBA").getchannel("A").getextrema()[0] < 255


def test_docx_exporter_uses_font_argument(tmp_path):
    output_path = tmp_path / "exported.docx"
    args = DefaultArgs(
        docx_template=os.path.join(parentdir, "chgksuite", "resources", "template.docx"),
        font="Test Embed",
        game="chgk",
        regexes_file=os.path.join(parentdir, "chgksuite", "resources", "regexes_ru.json"),
        optimize_size="off",
        spoilers="off",
        screen_mode="off",
    )

    exporter = DocxExporter(
        parse_4s("# Test\n\n? Question\n! Answer"),
        args,
        {"tmp_dir": str(tmp_path), "targetdir": str(tmp_path)},
    )
    exporter.export(output_path)

    with zipfile.ZipFile(output_path) as docx_file:
        assert not any(name.startswith("word/fonts/") for name in docx_file.namelist())
        font_table = docx_file.read("word/fontTable.xml").decode("utf-8")

    assert "Test Embed" in font_table


def test_docx_screen_mode_preserves_zachet_brackets(tmp_path):
    from docx import Document

    output_path = tmp_path / "screen.docx"
    args = DefaultArgs(
        docx_template=os.path.join(parentdir, "chgksuite", "resources", "template.docx"),
        game="chgk",
        regexes_file=os.path.join(parentdir, "chgksuite", "resources", "regexes_ru.json"),
        optimize_size="off",
        spoilers="off",
        screen_mode="replace_all",
    )
    exporter = DocxExporter(
        [
            (
                "Question",
                {
                    "question": "Вопрос [убрать].",
                    "answer": "Ответ [оставить].",
                    "zachet": "Зачет [оставить].",
                    "nezachet": "Незачет [убрать].",
                    "comment": "Комментарий [убрать].",
                },
            )
        ],
        args,
        {"tmp_dir": str(tmp_path), "targetdir": str(tmp_path)},
    )
    exporter.export(output_path)

    text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

    assert "Вопрос." in text
    assert "Ответ [оставить]." in text
    assert "Зачет [оставить]." in text
    assert "Незачет." in text
    assert "Комментарий." in text
    assert "[убрать]" not in text


def test_docx_cli_accepts_font():
    import argparse

    from chgksuite.cli import ArgparseBuilder

    parser = argparse.ArgumentParser(prog="chgksuite")
    ArgparseBuilder(parser, False).build()

    args = parser.parse_args(["compose", "docx", "--font", "Test Embed", "test.4s"])
    legacy_args = parser.parse_args(
        [
            "compose",
            "docx",
            "--font_face",
            "Legacy Font",
            "--optimize_size",
            "off",
            "test.4s",
        ]
    )

    assert args.font == "Test Embed"
    assert args.optimize_size == "on"
    assert legacy_args.font == "Legacy Font"
    assert legacy_args.optimize_size == "off"

    pptx_args = parser.parse_args(["compose", "pptx", "--font", "Test PPT", "test.4s"])
    pptx_no_optimize_args = parser.parse_args(
        ["compose", "pptx", "--optimize_size", "off", "test.4s"]
    )
    assert pptx_args.font == "Test PPT"
    assert pptx_args.optimize_size == "on"
    assert pptx_no_optimize_args.optimize_size == "off"


def test_telegram_formats_non_ascii_url_as_html_link():
    url = "https://ru.wikipedia.org/wiki/Абвгде_ёжзийкл_мнопрстуф_«Хцчшщы_—_эюяабв»_гдежзий_клмнопрст_уфхцчшщ?x=1&y=2"
    exporter = TelegramExporter.__new__(TelegramExporter)
    exporter.args = DefaultArgs()
    exporter.parse_4s_elem = _parse_4s_elem

    telegram_text, image = TelegramExporter.tgformat(exporter, url)

    assert telegram_text.startswith('<a href="https://ru.wikipedia.org/wiki/%D0%90')
    assert "%C2%AB%D0%A5%D1%86%D1%87%D1%88%D1%89%D1%8B" in telegram_text
    assert "%E2%80%94" in telegram_text
    assert "x=1&amp;y=2" in telegram_text
    assert "Абвгде&#95;ёжзийкл" in telegram_text
    assert telegram_text.endswith("</a>")
    assert image is None


def test_long_handout():
    cwd = os.getcwd()
    with make_temp_directory(dir=".") as temp_dir:
        shutil.copy(os.path.join(currentdir, "test.jpg"), temp_dir)
        shutil.copy(os.path.join(currentdir, "long_handout.png"), temp_dir)
        os.chdir(temp_dir)
        assert TelegramExporter.prepare_image_for_telegram("test.jpg") == "test.jpg"
        assert (
            TelegramExporter.prepare_image_for_telegram("long_handout.png")
            == "long_handout_telegram.jpg"
        )
        img = Image.open("long_handout_telegram.jpg")
        assert img.size == (1600, 83)
        os.chdir(cwd)


REPLACE_COUNTER_TEST_CASES = [
    ("4SCOUNTER 4SCOUNTER 4SCOUNTER", "1 2 3"),
    ("4SCOUNTER 4SCOUNTER1 4SCOUNTERa", "1 1 1"),
    ("set 4SCOUNTER = 5 4SCOUNTER", " 5"),
    ("set 4SCOUNTERa = 4 4SCOUNTERa", " 4"),
]


@pytest.mark.parametrize("replace_input, replace_output", REPLACE_COUNTER_TEST_CASES)
def test_replace_counters(replace_input, replace_output):
    assert replace_counters(replace_input) == replace_output
